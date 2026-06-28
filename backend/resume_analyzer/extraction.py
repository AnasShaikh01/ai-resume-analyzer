import io
import re
import os

import docx
import fitz  # PyMuPDF
import pandas as pd

from .skill_extractor import SkillExtractor
from .resume_section_parser import ResumeSectionParser

# ==========================================================
# Skill Loading
# ==========================================================

def load_skills(csv_path="skills.csv"):
    """
    Loads skills from CSV.

    Supported formats:

    1 Column:
        Python
        Java
        SQL

    2 Columns:
        python,py
        javascript,js

    Returns:
        list OR dict
    """

    current_dir = os.path.dirname(os.path.abspath(__file__))

    absolute_csv_path = os.path.join(
        current_dir,
        "..",
        "..",
        csv_path
    )

    try:
        df = pd.read_csv(absolute_csv_path, header=None)

    except FileNotFoundError:
        return []

    # Single-column skills
    if df.shape[1] == 1:

        return [
            str(skill).strip().lower()
            for skill in df[0].dropna().tolist()
        ]

    # Canonical + aliases
    elif df.shape[1] >= 2:

        skills_dict = {}

        for _, row in df.iterrows():

            canonical = str(row[0]).strip().lower()

            alias = ""

            if pd.notna(row[1]):
                alias = str(row[1]).strip().lower()

            if canonical:

                if canonical not in skills_dict:
                    skills_dict[canonical] = []

                if alias and alias not in skills_dict[canonical]:
                    skills_dict[canonical].append(alias)

        return skills_dict

    return []


# ==========================================================
# NLP Skill Extraction
# ==========================================================

_skill_extractor_cache = {}
_section_parser = ResumeSectionParser()


def extract_skills(text, skills_source):

    cache_key = str(id(skills_source))
    if cache_key not in _skill_extractor_cache:
        _skill_extractor_cache[cache_key] = SkillExtractor(
            skills_source
        )
    return _skill_extractor_cache[
        cache_key
    ].extract(text)

def extract_skills_with_context(text, skills_source):

    cache_key = str(id(skills_source))
    if cache_key not in _skill_extractor_cache:
        _skill_extractor_cache[cache_key] = SkillExtractor(
            skills_source
        )
    return _skill_extractor_cache[
        cache_key
    ].extract_with_context(text)

def parse_resume_sections(text):
    """
    Parse resume into sections.

    Returns:
    {
        "skills": "...",
        "projects": "...",
        "experience": "...",
        ...
    }
    """

    return _section_parser.parse(text)


# ==========================================================
# ATS Utilities
# ==========================================================

def _normalize_font_name(name: str) -> str:
    """
    Example:
    AAAAAA+Calibri-Bold -> Calibri
    """
    if not name:
        return ""
    name = name.split("+")[-1]
    name = re.sub(
        r"[- ]?(Bold|Italic|Oblique|Regular)$",
        "",
        name,
        flags=re.IGNORECASE
    )
    return name.strip()


def _detect_multi_column(page) -> bool:
    """
    Detect if PDF page contains multiple columns.
    """
    blocks = page.get_text("blocks")
    if not blocks or len(blocks) < 10:
        return False
    x_positions = [b[0] for b in blocks]
    clustered = [
        round(x / 20) * 20
        for x in x_positions
    ]
    cluster_counts = {}

    for c in clustered:
        cluster_counts[c] = cluster_counts.get(c, 0) + 1
    sorted_clusters = sorted(
        cluster_counts.items(),
        key=lambda x: -x[1]
    )

    if len(sorted_clusters) >= 2:
        (x1, count1), (x2, count2) = sorted_clusters[:2]
        distance = abs(x1 - x2)
        if (
            distance > page.rect.width * 0.35
            and count1 > 5
            and count2 > 5
        ):
            return True

    return False

def check_ats_compliance(file_bytes, file_type: str):
    """
    Check ATS-unfriendly formatting.
    """
    results = {
        "multi_column": False,
        "non_standard_fonts": False,
        "images": False,
        "tables": False
    }

    bad_fonts = {
        "Comic Sans",
        "Papyrus",
        "Brush Script",
        "Cursive",
        "Monotype Corsiva"
    }

    try:
        if file_type == "pdf":
            doc = fitz.open(
                stream=file_bytes,
                filetype="pdf"
            )
            for page in doc:
                # Detect images
                for img in page.get_images(full=True):
                    xref = img[0]
                    pix = fitz.Pixmap(doc, xref)
                    if (
                        pix.n >= 3
                        and pix.width > 100
                        and pix.height > 100
                    ):
                        results["images"] = True
                        break
                # Detect non-standard fonts
                for f in page.get_fonts():
                    font_name = _normalize_font_name(f[3])
                    if any(
                        bad.lower() in font_name.lower()
                        for bad in bad_fonts
                    ):
                        results["non_standard_fonts"] = True
                        break
                # Detect multiple columns
                if _detect_multi_column(page):
                    results["multi_column"] = True
                if any(results.values()):
                    break
            doc.close()

        elif file_type == "docx":
            doc = docx.Document(
                io.BytesIO(file_bytes)
            )
            if len(doc.tables) > 0:
                results["tables"] = True
            if len(doc.inline_shapes) > 0:
                results["images"] = True
            if (
                len(doc.sections) > 1
                and doc.sections[0].start_type != 2
            ):
                results["multi_column"] = True

            for para in doc.paragraphs:
                for run in para.runs:
                    if (
                        run.font.name
                        and any(
                            bad.lower() in run.font.name.lower()
                            for bad in bad_fonts
                        )
                    ):
                        results["non_standard_fonts"] = True
                        break
                if results["non_standard_fonts"]:
                    break
    except Exception as e:
        results["error"] = str(e)
    return results